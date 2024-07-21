from docx import Document

# Load the DOCX file
doc = Document("Description Demo.docx")

# Apply corrections and improvements
for paragraph in doc.paragraphs:
    if "bootstap" in paragraph.text:
        # Correct the typo from 'bootstap' to 'bootstrap'
        paragraph.text = paragraph.text.replace("bootstap", "bootstrap")

# Add a header to the document
doc.add_heading('Description Demo - Improved Version', level=1)

# Save the corrected version
corrected_path = "/mnt/data/Description Demo Improved.docx"
doc.save(corrected_path)

# Convert to PDF using python-docx (Requires additional library)
from fpdf import FPDF

class PDF(FPDF):
    def header(self):
        self.set_font('Arial', 'B', 12)
        self.cell(0, 10, 'Description Demo - Improved Version', 0, 1, 'C')

    def footer(self):
        self.set_y(-15)
        self.set_font('Arial', 'I', 8)
        self.cell(0, 10, f'Page {self.page_no()}', 0, 0, 'C')

# Convert DOCX text to PDF
pdf = PDF()
pdf.add_page()

doc = Document(corrected_path)
for para in doc.paragraphs:
    pdf.set_font('Arial', '', 12)
    pdf.multi_cell(0, 10, para.text)

pdf_output = "Description Demo Improved.pdf"
pdf.output(pdf_output)
pdf_output
